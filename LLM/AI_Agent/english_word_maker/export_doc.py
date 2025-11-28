import json
from io import BytesIO

import pandas as pd
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

KOREAN_FONT_PATH = r"C:\Windows\Fonts\malgun.ttf"


def get_excel_bytes(df: pd.DataFrame):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Vocab")
    output.seek(0)
    return output


def clean_text(value) -> str:
    if value is None:
        return ""
    text = str(value)
    text = text.replace("\r\n", " ").replace("\n", " ").replace("\t", " ")
    text = "".join(ch if ord(ch) >= 32 else " " for ch in text)
    return text


def render_entry_to_doc(doc: Document, entry: dict):
    word = entry.get("word", "")
    pos = entry.get("pos", "")
    ipa = entry.get("ipa", "")
    lookup = entry.get("lookup_count", 1)

    p = doc.add_paragraph()
    run = p.add_run(f"{word} ")
    run.bold = True
    p.add_run(f"({pos}) {ipa}  [lookup: {lookup}]")

    domain = entry.get("domain", "")
    level = entry.get("level", "")
    freq = entry.get("frequency", "")
    etym = entry.get("etymology", "")
    similars = entry.get("similar_words", [])

    doc.add_paragraph(f"Domain: {domain}  |  Level: {level}  |  Freq: {freq}")
    if etym:
        doc.add_paragraph(f"Etymology: {etym}")
    if similars:
        doc.add_paragraph("Similar: " + ", ".join(similars))

    senses_ko = entry.get("senses_ko", {})
    if isinstance(senses_ko, dict) and senses_ko:
        doc.add_paragraph("[Meaning-KO]")
        for pos_key, m_ko in senses_ko.items():
            doc.add_paragraph(f"({pos_key}) {m_ko}")

    senses_en = entry.get("senses_en", {})
    if isinstance(senses_en, dict) and senses_en:
        doc.add_paragraph("[Meaning-EN]")
        for pos_key, m_en in senses_en.items():
            doc.add_paragraph(f"({pos_key}) {m_en}")

    examples = entry.get("examples", {})
    if isinstance(examples, dict) and examples:
        doc.add_paragraph("[Examples]")
        for pos_key, ex in examples.items():
            doc.add_paragraph(f"({pos_key}) {ex}")

    syns = entry.get("synonyms", [])
    if syns:
        doc.add_paragraph("Synonyms: " + ", ".join(syns[:3]))

    ants = entry.get("antonyms", [])
    if ants:
        doc.add_paragraph("Antonyms: " + ", ".join(ants[:3]))

    doc.add_paragraph("-" * 50)


def get_word_bytes_from_vocab(vocab: dict):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    doc.add_heading("단어장", level=1)

    for _, entry in vocab.items():
        render_entry_to_doc(doc, entry)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def pdf_write_line(pdf: FPDF, text: str, max_chars: int = 80):
    text = clean_text(text)
    while text:
        chunk = text[:max_chars]
        pdf.cell(0, 8, chunk, ln=1)
        text = text[max_chars:]


def render_entry_to_pdf(pdf: FPDF, entry: dict):
    word = entry.get("word", "")
    pos = entry.get("pos", "")
    ipa = entry.get("ipa", "")
    lookup = entry.get("lookup_count", 1)

    header = f"{word} ({pos}) {ipa}  [lookup: {lookup}]"
    pdf_write_line(pdf, header)

    domain = entry.get("domain", "")
    level = entry.get("level", "")
    freq = entry.get("frequency", "")
    similars = entry.get("similar_words", [])

    meta_line = f"Domain: {domain}  |  Level: {level}  |  Freq: {freq}"
    pdf_write_line(pdf, meta_line)
    if similars:
        pdf_write_line(pdf, "Similar: " + ", ".join(similars))

    senses_ko = entry.get("senses_ko", {})
    if isinstance(senses_ko, dict) and senses_ko:
        pdf_write_line(pdf, "[Meaning-KO]")
        for pos_key, m_ko in senses_ko.items():
            pdf_write_line(pdf, f"({pos_key}) {m_ko}")

    senses_en = entry.get("senses_en", {})
    if isinstance(senses_en, dict) and senses_en:
        pdf_write_line(pdf, "[Meaning-EN]")
        for pos_key, m_en in senses_en.items():
            pdf_write_line(pdf, f"({pos_key}) {m_en}")

    examples = entry.get("examples", {})
    if isinstance(examples, dict) and examples:
        pdf_write_line(pdf, "[Examples]")
        for pos_key, ex in examples.items():
            pdf_write_line(pdf, f"({pos_key}) {ex}")

    syns = entry.get("synonyms", [])
    if syns:
        pdf_write_line(pdf, "Synonyms: " + ", ".join(syns[:3]))

    ants = entry.get("antonyms", [])
    if ants:
        pdf_write_line(pdf, "Antonyms: " + ", ".join(ants[:3]))

    pdf_write_line(pdf, "-" * 50)
    pdf.ln(4)


def get_pdf_bytes_from_vocab(vocab: dict):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.add_font("malgun", "", KOREAN_FONT_PATH, uni=True)
    pdf.set_font("malgun", size=12)

    for _, entry in vocab.items():
        render_entry_to_pdf(pdf, entry)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf
