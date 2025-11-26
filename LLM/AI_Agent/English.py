# web
import streamlit as st
from langchain_core.messages.chat import ChatMessage

# LLM
from langchain_core.prompts import ChatPromptTemplate  # prompt
from langchain_core.prompts import load_prompt  # prompt load

from langchain import hub  # prompt hub
from langchain_openai import ChatOpenAI  # AI
from langchain_core.output_parsers import StrOutputParser  # 출력
from dotenv import load_dotenv  # 설정 값

# 단어장 다운로드
import json, re
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from pathlib import Path  # 단어장 위치

from langchain_teddynote import logging  # langsmith_logging

# 단어장 발음 다운로드(251126)
from openai import OpenAI
from zipfile import ZipFile

logging.langsmith("English_word_maker_project")

load_dotenv()

st.title("my chatgpt")

# 단어장 저장 위치
SAVE_DIR = Path("English_word")
SAVE_PATH = SAVE_DIR / "vocab.json"


# 처음 1번만 실행하기 위한 코드
if "messages" not in st.session_state:
    # 대화기록을 저장하기 위해 생성
    st.session_state["messages"] = []

# 단어장 생성을 위한 초기 코드
if "vocab" not in st.session_state:
    st.session_state["vocab"] = {}
    if SAVE_PATH.exists():
        try:
            with open(SAVE_PATH, "r", encoding="utf-8") as f:
                loaded = json.load(f)
            # loaded 는 {word_key: entry_dict} 형태라고 가정
            st.session_state["vocab"] = loaded
        except Exception as e:
            st.warning(f"저장된 단어장 로드 실패: {e}")

# 사이드바 생성
with st.sidebar:
    # 대화 초기화 버튼
    clear_bar = st.button("대화 초기화")

    option = st.selectbox(
        "Please Select Prompt", ("Basic", "English", "Summary"), index=0
    )


# 이전 대화 출력
def print_messages():
    for chat_message in st.session_state["messages"]:
        st.chat_message(chat_message.role).write(chat_message.content)


# 대화 초기화
if clear_bar:
    st.session_state["messages"] = []
# 이전 대화 기록 출력
print_messages()


# 새로운 메시지 추가
def add_message(role, message):
    st.session_state["messages"].append(ChatMessage(role=role, content=message))


# chain 생성
def create_chain(prompt_type):
    # 프롬프트(기본 모드)
    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", "당신은 친절한 AI 어시스턴트입니다. 자세히 답변해 주세요."),
            ("user", "#Question:\n{question}"),
        ]
    )
    if prompt_type == "English":
        prompt = load_prompt("prompts/English.yaml", encoding="utf-8")
        prompt
    elif prompt_type == "Summary":
        prompt = hub.pull("baektest/chain-of-density-map-korean", include_model=True)
        # prompt = hub.pull("teddynote/chain-of-density-prompt", include_model=True)

    # AI
    llm = ChatOpenAI(model_name="gpt-5.1", temperature=0)  # gpt-5 오류 발생
    # 출력
    output_parser = StrOutputParser()

    # 체인 생성
    chain = prompt | llm | output_parser
    return chain


# 품사 정의
POS_MAP = {
    "n": "noun",
    "v": "verb",
    "adj": "adjective",
    "adv": "adverb",
}


# 한글 의미 품사별 dict로 변경
def extract_senses_ko(meaning_text: str) -> dict:
    if not meaning_text:
        return {}

    senses_ko = {}
    lines = [l.strip() for l in meaning_text.splitlines() if l.strip()]

    for line in lines:
        m = re.match(r"^\s*[-•*]?\s*([a-zA-Z]+)[\)\.\:]\s*(.+)", line)
        if not m:
            continue

        short_pos = m.group(1).lower()
        ko_mean = m.group(2).strip()
        pos_full = POS_MAP.get(short_pos, short_pos)

        senses_ko[pos_full] = ko_mean

    return senses_ko


# 영어 의미 품사별 dict로 변경
def extract_senses_en(senses_text: str) -> dict:
    if not senses_text:
        return {}

    senses_en = {}
    lines = [l.strip() for l in senses_text.splitlines() if l.strip()]

    for line in lines:
        m = re.match(r"^\s*[-•*]?\s*([a-zA-Z]+)[\)\.\:]\s*(.+)", line)
        if not m:
            continue

        short_pos = m.group(1).lower()
        en_mean = m.group(2).strip()
        pos_full = POS_MAP.get(short_pos, short_pos)

        senses_en[pos_full] = en_mean

    return senses_en


# 유사단어가 다수 있을 경우 ',' 문자열을 리스트로 변경
def extract_similar_words(confusable_text: str):
    if not confusable_text:
        return []
    return [x.strip() for x in confusable_text.split(",") if x.strip()]


# 프롬프트 설정에 따른 단어장 생성
def parse_yaml_style_answer(answer_text: str) -> dict:
    # 번호 -> 내부 key 매핑
    field_map = {
        1: "word",
        2: "domain",
        3: "pos",
        4: "ipa",
        5: "senses_ko_raw",
        6: "senses_en_raw",
        7: "level",
        8: "frequency",
        9: "etymology",
        10: "examples_raw",
        11: "synonyms",
        12: "antonyms",
        13: "collocations",
        14: "derivatives",
        15: "tags",
        16: "confusable",
        17: "lookup_count_raw",  # LLM이 적어준 값(우리는 참고만)
    }

    # "1) ", "2) " ... 로 시작하는 부분 찾기
    pattern = re.compile(r"^\s*(\d+)\)\s", re.MULTILINE)
    matches = list(pattern.finditer(answer_text))

    result: dict[str, object] = {}

    if not matches:
        return result  # 파싱 실패 시 빈 dict

    for idx, m in enumerate(matches):
        num = int(m.group(1))
        start = m.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(answer_text)

        block = answer_text[start:end].strip()

        # "N) 제목(Title): 내용..." 형태에서 첫 번째 ":" 뒤를 값으로 사용
        header_split = block.split(":", 1)
        if len(header_split) == 2:
            value = header_split[1].strip()
        else:
            lines = block.splitlines()
            if len(lines) >= 2:
                value = "\n".join(lines[1:]).strip()
            else:
                value = ""

        key = field_map.get(num)
        if key:
            result[key] = value

    # word가 여러 줄일 수 있으니 첫 줄만 사용
    if "word" in result and result["word"]:
        first_line = result["word"].splitlines()[0].strip()
        result["word"] = first_line

    # 태그 / 유의어 / 반의어 / 콜로케이션 등은 쉼표 기준으로 구분된 필드 설정
    for list_key in ["synonyms", "antonyms", "collocations", "derivatives", "tags"]:
        if list_key in result and result[list_key]:
            # "a, b, c" -> ["a", "b", "c"]
            items = [x.strip() for x in result[list_key].split(",") if x.strip()]
            result[list_key] = items
        else:
            result[list_key] = items

    # 10) 예문 JSON 파싱 처리(내용 추가 25.11.25)
    examples = {}
    if "examples_raw" in result and result["examples_raw"]:
        ex_raw = result["examples_raw"]

        # JSON 부분만 추출: { ... } 안
        json_match = re.search(r"\{[\s\S]*\}", ex_raw)
        if json_match:
            json_text = json_match.group(0)
            try:
                # JSON을 dict로 파싱
                examples = json.loads(json_text)
            except Exception:
                # 파싱 안 되면 그나마 깔끔하게 정리
                examples = {"raw": ex_raw}
        else:
            # 다른 형식에 따른 작성
            temp = {}
            for line in ex_raw.splitlines():
                line = line.strip()
                if not line:
                    continue
                # pos: sentence...
                if ":" in line:
                    k, v = line.split(":", 1)
                    temp[k.strip()] = v.strip()
            examples = temp if temp else {"raw": ex_raw}
    result["examples"] = examples

    # 5) 뜻(Meaning) → senses_ko (품사별 한글 의미) 추가
    senses_ko_raw = str(result.get("senses_ko_raw", "") or "")
    result["senses_ko"] = extract_senses_ko(senses_ko_raw)

    # 6) 영어 의미(Senses_En) → senses_en (품사별 영어 의미) 추가
    senses_en_raw = str(result.get("senses_en_raw", "") or "")
    result["senses_en"] = extract_senses_en(senses_en_raw)

    # 16) 유사단어(Confusable) → similar_words 리스트 추가
    confusable_text = str(result.get("confusable", "") or "")
    result["similar_words"] = extract_similar_words(confusable_text)

    return result


# 단어장 생성용 함수
def update_vocab_from_answer(answer_text: str):
    # LLM이 넘겨준 JSON 텍스트(answer_text)를 파싱해서 단어장에 반영
    parsed = parse_yaml_style_answer(answer_text)
    word = parsed.get("word", "")

    if not word:
        return  # 단어 못 찾으면 무시

    word_key = word.lower().strip()
    vocab = st.session_state["vocab"]

    if word_key in vocab:
        # 이미 있는 단어 → 조회수 +1
        vocab[word_key]["lookup_count"] += 1

    else:
        # 새 단어 → parsed 정보를 그대로 넣고 lookup_count는 1부터 시작
        entry = parsed.copy()
        entry["lookup_count"] = 1  # LLM이 써준 15번 필드는 무시하고 우리가 관리
        vocab[word_key] = entry

    return word_key  # 최근 단어 저장용


# 단어장을 저장하기 위한 작업(DF화)
def vocab_to_df() -> pd.DataFrame:
    vocab = st.session_state.get("vocab", {})
    rows = []

    for word_key, info in vocab.items():
        senses_ko = info.get("senses_ko", {})
        senses_en = info.get("senses_en", {})
        examples = info.get("examples", {})
        synonyms = info.get("synonyms", [])
        antonyms = info.get("antonyms", [])
        collocs = info.get("collocations", [])
        derivatives = info.get("derivatives", [])
        tags = info.get("tags", [])
        similar_words = info.get("similar_words", [])
        confusable_raw = info.get("confusable", "")

        rows.append(
            {
                "Word": info.get("word", word_key),
                "Domain": info.get("domain", ""),
                "POS": info.get("pos", ""),
                "IPA": info.get("ipa", ""),
                "Level": info.get("level", ""),
                "Frequency": info.get("frequency", ""),
                "Etymology": info.get("etymology", ""),
                "LookupCount": info.get("lookup_count", 1),
                "Senses_Ko": json.dumps(senses_ko, ensure_ascii=False),
                "Senses_En": json.dumps(senses_en, ensure_ascii=False),
                "Examples": json.dumps(examples, ensure_ascii=False),
                "Synonyms": json.dumps(synonyms, ensure_ascii=False),
                "Antonyms": json.dumps(antonyms, ensure_ascii=False),
                "Collocations": json.dumps(collocs, ensure_ascii=False),
                "Derivatives": json.dumps(derivatives, ensure_ascii=False),
                "Tags": json.dumps(tags, ensure_ascii=False),
                "SimilarWords": json.dumps(similar_words, ensure_ascii=False),
                "ConfusableRaw": confusable_raw,
            }
        )

    if rows:
        return pd.DataFrame(rows)
    else:
        return pd.DataFrame(
            columns=[
                "Word",
                "Domain",
                "POS",
                "IPA",
                "Level",
                "Frequency",
                "Etymology",
                "LookupCount",
                "Senses_Ko",
                "Senses_En",
                "Examples",
                "Synonyms",
                "Antonyms",
                "Collocations",
                "Derivatives",
                "Tags",
                "SimilarWords",
                "ConfusableRaw",
            ]
        )


# Excel
def get_excel_bytes(df: pd.DataFrame):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Vocab")
    output.seek(0)
    return output


# Word(docx)
def get_word_bytes(df: pd.DataFrame):
    doc = Document()
    doc.add_heading("단어장", level=1)

    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# PDF 용 text 정리 함수
def clean_text(value) -> str:
    """fpdf2에 넣기 전에 텍스트를 안전한 형태로 정리"""
    if value is None:
        return ""
    text = str(value)

    # 줄바꿈/캐리지리턴/탭 같은 제어 문자 정리
    text = text.replace("\r\n", " ")
    text = text.replace("\n", " ")
    text = text.replace("\t", " ")

    # fpdf가 싫어할 수 있는 control char 제거 (ASCII 32 미만)
    text = "".join(ch if ord(ch) >= 32 else " " for ch in text)

    return text


KOREAN_FONT_PATH = r"C:\Windows\Fonts\malgun.ttf"


# PDF
"""
def get_pdf_bytes(df):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # 1 유니코드 TTF 폰트 등록 + 사용
    pdf.add_font("malgun", "", KOREAN_FONT_PATH, uni=True)
    pdf.set_font("malgun", size=12)

    # 2 여기서부터는 아무 텍스트나(한글/영문/특수기호) 써도 됨
    for _, row in df.iterrows():
        # 모든 텍스트를 한 번 정리
        word = clean_text(row.get("Word", ""))
        pos = clean_text(row.get("POS", ""))
        lookup = clean_text(row.get("LookupCount", ""))

        meaning_src = row.get("Meaning", "")
        meaning = clean_text(meaning_src)

        # 1) 첫 줄: multi_cell 대신 cell 사용 (줄바꿈 알고리즘 안 탄다)
        line1 = f"{word} ({pos})  [lookup {lookup}]"
        safe_line1 = line1[:200]  # 혹시 모를 경우를 위해 너무 길면 잘라줌
        pdf.cell(0, 8, safe_line1, ln=1)  # ln=1 → 이 줄 출력 후 줄바꿈

        # 2) 의미: 길 수 있으니 우리가 직접 잘라서 여러 줄로 찍기
        prefix = "- mean: "
        text = prefix + meaning
        max_chars = 80  # 한 줄에 찍을 최대 문자 수 (대략)

        while text:
            chunk = text[:max_chars]
            pdf.cell(0, 8, chunk, ln=1)
            text = text[max_chars:]

        # ===== 3) 예시 문장 (최대 2개) =====
        examples_raw = str(row.get("Examples", "") or "")
        examples_raw = clean_text(examples_raw)
        if examples_raw:
            pdf.cell(0, 8, "- examples:", ln=1)
            # 줄 단위로 쪼개서 앞 2개만 사용
            example_lines = [e.strip() for e in examples_raw.split("\n") if e.strip()]
            for ex in example_lines[:2]:
                ex_line = f"  · {ex}"
                # 너무 길면 자르기
                while ex_line:
                    chunk = ex_line[:max_chars]
                    pdf.cell(0, 8, chunk, ln=1)
                    ex_line = ex_line[max_chars:]

        # ===== 4) 유의어 (최대 3개) =====
        syn_raw = str(row.get("Synonyms", "") or "")
        syn_list = [s.strip() for s in syn_raw.split(",") if s.strip()]
        syn_list = syn_list[:3]  # 최대 3개
        if syn_list:
            syn_line = "- synonyms: " + ", ".join(syn_list)
            syn_line = clean_text(syn_line)
            while syn_line:
                chunk = syn_line[:max_chars]
                pdf.cell(0, 8, chunk, ln=1)
                syn_line = syn_line[max_chars:]

        # ===== 5) 반의어 (최대 3개) =====
        ant_raw = str(row.get("Antonyms", "") or "")
        ant_list = [a.strip() for a in ant_raw.split(",") if a.strip()]
        ant_list = ant_list[:3]
        if ant_list:
            ant_line = "- antonyms: " + ", ".join(ant_list)
            ant_line = clean_text(ant_line)
            while ant_line:
                chunk = ant_line[:max_chars]
                pdf.cell(0, 8, chunk, ln=1)
                ant_line = ant_line[max_chars:]

        # 단어 사이에 빈 줄 하나
        pdf.ln(4)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf
"""


# 단어장 저장 함수
def save_vocab_to_disk():
    """현재 session_state['vocab']를 JSON으로 저장"""
    try:
        SAVE_DIR.mkdir(exist_ok=True)
        with open(SAVE_PATH, "w", encoding="utf-8") as f:
            json.dump(
                st.session_state.get("vocab", {}),
                f,
                ensure_ascii=False,
                indent=2,
            )
    except Exception as e:
        st.error(f"단어장 저장 중 오류 발생: {e}")


# Word 포맷 출력
def render_entry_to_doc(doc: Document, entry: dict):
    # 1. 단어(품사) / 발음 기호 / 문의 횟수
    word = entry.get("word", "")
    pos = entry.get("pos", "")
    ipa = entry.get("ipa", "")
    lookup = entry.get("lookup_count", 1)

    p = doc.add_paragraph()
    run = p.add_run(f"{word} ")
    run.bold = True
    p.add_run(f"({pos}) {ipa}  [lookup: {lookup}]")

    # 2. 분야 / 난이도 / 빈출도 / 어원 / 유사단어
    domain = entry.get("domain", "")
    level = entry.get("level", "")
    freq = entry.get("frequency", "")
    etym = entry.get("etymology", "")
    similars = entry.get("similar_words", [])

    doc.add_paragraph(f"Domain: {domain}  |  Level: {level}  |  Freq: {freq}")
    if etym:
        doc.add_paragraph(f"Etymology: {etym}")
    if similars:
        doc.add_paragraph("Similar: " + ", ".join(similars))

    # 3. 품사별 의미(한글)
    senses_ko = entry.get("senses_ko", {})
    if isinstance(senses_ko, dict) and senses_ko:
        doc.add_paragraph("[Meaning-KO]")
        for pos_key, m_ko in senses_ko.items():
            doc.add_paragraph(f"({pos_key}) {m_ko}")

    # 4. 품사별 의미(영어)
    senses_en = entry.get("senses_en", {})
    if isinstance(senses_en, dict) and senses_en:
        doc.add_paragraph("[Meaning-EN]")
        for pos_key, m_en in senses_en.items():
            doc.add_paragraph(f"({pos_key}) {m_en}")

    # 5. 예문 품사별 1건
    examples = entry.get("examples", {})
    if isinstance(examples, dict) and examples:
        doc.add_paragraph("[Examples]")
        for pos_key, ex in examples.items():
            doc.add_paragraph(f"({pos_key}) {ex}")

    # 6. 유의어(최대 3)
    syns = entry.get("synonyms", [])
    if syns:
        doc.add_paragraph("Synonyms: " + ", ".join(syns[:3]))

    # 7. 반의어(최대 3)
    ants = entry.get("antonyms", [])
    if ants:
        doc.add_paragraph("Antonyms: " + ", ".join(ants[:3]))

    # 8. 구분선
    doc.add_paragraph("-" * 50)


# Word 포맷에 따른 수정
def get_word_bytes_from_vocab(vocab: dict):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    doc.add_heading("단어장", level=1)

    for word_key, entry in vocab.items():
        render_entry_to_doc(doc, entry)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# PDF 포맷 변경에 따른 수정(1줄 출력)
def pdf_write_line(pdf: FPDF, text: str, max_chars: int = 80):
    text = clean_text(text)
    while text:
        chunk = text[:max_chars]
        pdf.cell(0, 8, chunk, ln=1)
        text = text[max_chars:]


# PDF 포맷 설정
def render_entry_to_pdf(pdf: FPDF, entry: dict):
    word = entry.get("word", "")
    pos = entry.get("pos", "")
    ipa = entry.get("ipa", "")
    lookup = entry.get("lookup_count", 1)

    # 1. 단어(품사) / 발음 기호 / 문의 횟수
    header = f"{word} ({pos}) {ipa}  [lookup: {lookup}]"
    pdf_write_line(pdf, header)

    # 2. 분야 / 난이도 / 빈출도 / 유사단어
    domain = entry.get("domain", "")
    level = entry.get("level", "")
    freq = entry.get("frequency", "")
    similars = entry.get("similar_words", [])

    meta_line = f"Domain: {domain}  |  Level: {level}  |  Freq: {freq}"
    pdf_write_line(pdf, meta_line)
    if similars:
        pdf_write_line(pdf, "Similar: " + ", ".join(similars))

    # 3. 품사별 의미(한글)
    senses_ko = entry.get("senses_ko", {})
    if isinstance(senses_ko, dict) and senses_ko:
        pdf_write_line(pdf, "[Meaning-KO]")
        for pos_key, m_ko in senses_ko.items():
            pdf_write_line(pdf, f"({pos_key}) {m_ko}")

    # 4. 품사별 의미(영어)
    senses_en = entry.get("senses_en", {})
    if isinstance(senses_en, dict) and senses_en:
        pdf_write_line(pdf, "[Meaning-EN]")
        for pos_key, m_en in senses_en.items():
            pdf_write_line(pdf, f"({pos_key}) {m_en}")

    # 5. 예문 품사별 1건
    """
    examples = entry.get("examples", {})
    if isinstance(examples, dict) and examples:
        pdf_write_line(pdf, "[Examples]")
        for pos_key, ex in examples.items():
            pdf_write_line(pdf, f"({pos_key}) {ex}")
    """
    # 5-1. 공통함수 사용(예문 발음/pdf)
    examples_for_pdf = get_examples_for_pdf(entry)
    if examples_for_pdf:
        pdf_write_line(pdf, "[Examples]")
        for pos_key, ex in examples_for_pdf:
            pdf_write_line(pdf, f"({pos_key}) {ex}")

    # 6. 유의어(최대 3)
    syns = entry.get("synonyms", [])
    if syns:
        pdf_write_line(pdf, "Synonyms: " + ", ".join(syns[:3]))

    # 7. 반의어(최대 3)
    ants = entry.get("antonyms", [])
    if ants:
        pdf_write_line(pdf, "Antonyms: " + ", ".join(ants[:3]))

    # 8. 구분선
    pdf_write_line(pdf, "-" * 50)
    pdf.ln(4)


# PDF 포맷 변경에 따른 수정
def get_pdf_bytes_from_vocab(vocab: dict):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.add_font("malgun", "", KOREAN_FONT_PATH, uni=True)
    pdf.set_font("malgun", size=12)

    for word_key, entry in vocab.items():
        render_entry_to_pdf(pdf, entry)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


# 단어/예문 발음 추가 - 25.11.26
client = OpenAI()

# 최근 조회 단어 확인
if "last_word_key" not in st.session_state:
    st.session_state["last_word_key"] = None


# TTS 함수
def tts_bytes(text: str) -> bytes:
    if not text:
        return b""

    try:
        response = client.audio.speech.create(
            model="tts-1",
            voice="coral",  # allay / coral / nova
            input=text,
        )
        audio_bytes = response.read()
        return audio_bytes
    except Exception as e:
        print("TTS error: ", e)
        return b""


# 음성파일 이름용 함수
def sanitize_filename(name: str) -> str:
    return re.sub(r"[^0-9a-zA-Z가-힣_]+", "_", name)


# 단어/예문 발음 zip으로
def get_audio_zip_from_vocab(vocab: dict) -> BytesIO:
    buf = BytesIO()
    with ZipFile(buf, "w") as z:
        for word_key, entry in vocab.items():
            word = entry.get("word", word_key)
            safe_word = sanitize_filename(word)

            # 1) 단어 발음
            try:
                word_audio = tts_bytes(word)
                if word_audio:
                    z.writestr(f"Pronunciation/{safe_word}.mp3", word_audio)
            except Exception as e:
                print("word tts error", word, e)

            # 2) 예문 발음 (각 품사별 1개씩)
            examples_for_pdf = get_examples_for_pdf(entry)
            for idx, (pos_key, ex) in enumerate(examples_for_pdf, start=1):
                try:
                    ex_audio = tts_bytes(ex)
                    if ex_audio:
                        z.writestr(
                            f"{safe_word}/{safe_word}_example{idx}.mp3",
                            ex_audio,
                        )
                except Exception as e:
                    print("example tts error:", word, pos_key, e)
    buf.seek(0)
    return buf


# pdf에 표시할 예문을 공통으로 가지고 오는 함수(예문 발음)
def get_examples_for_pdf(entry: dict):
    examples = entry.get("examples", {})
    if not isinstance(examples, dict) or not examples:
        return []
    result = []
    for pos_key, ex in examples.items():
        ex = (ex or "").strip()
        if not ex:
            continue
        result.append((pos_key, ex))

    return result


# 사용자 입력
user_input = st.chat_input("내용 입력")

# 입력 시
if user_input:
    # web 사용자 대화 출력(사용자 입력)
    st.chat_message("user").write(user_input)

    # 사용자 대화 기록 저장
    add_message("user", user_input)

    # chain 생성
    chain = create_chain(option)
    if option == "English":
        # 단어장 agent: 전체 응답을 한번에 받아서 처리
        answer_text = chain.invoke({"word": user_input})
        # 화면 출력
        st.chat_message("assistant").markdown(answer_text)
        add_message("assistant", answer_text)
        # 단어장 업데이트
        word_key = update_vocab_from_answer(answer_text)
        # 단어 발음 생성
        if word_key:
            st.session_state["last_word_key"] = word_key
    else:
        responce = chain.stream({user_input})
        # web assistant 대화 출력 방법2
        with st.chat_message("assistant"):
            # 빈 공간 생성
            container = st.empty()
            ai_answer = ""
            for token in responce:
                ai_answer += token
                container.markdown(ai_answer)

        # assistant 대화 저장
        add_message("assistant", ai_answer)

# 최근 조회한 단어 발음(UI)
st.subheader("최근 조회한 단어 발음")
last_key = st.session_state.get("last_word_key")
vocab = st.session_state.get("vocab", {})

if last_key and last_key in vocab:
    entry = vocab[last_key]
    word = entry.get("word", last_key)
    examples = entry.get("examples", {})

    st.write(f"**{word}** 발음")

    # 단어 발음 버튼
    if st.button("🔊 단어 발음 듣기"):
        audio = tts_bytes(word)
        if audio:
            st.audio(audio, format="audio/mp3")

    # 예문 발음 버튼 (최대 3개)
    if isinstance(examples, dict) and examples:
        st.write("예문 발음")
        for idx, (pos_key, ex) in enumerate(examples.items()):
            if idx >= 3:
                break
            col1, col2 = st.columns([3, 1])
            with col1:
                st.write(f"({pos_key}) {ex}")
            with col2:
                if st.button(f"🔊 예문 {idx+1}", key=f"ex_{last_key}_{idx}"):
                    audio = tts_bytes(ex)
                    if audio:
                        st.audio(audio, format="audio/mp3")
else:
    st.info("아직 조회한 단어가 없습니다.")

st.subheader("현재 단어장")
df = vocab_to_df()
add_message("assistant", df)
st.dataframe(df, use_container_width=True)

# Excel
if not df.empty:
    # 공통: 버튼 누르면 서버에 단어장 Json 저장
    save_vocab_to_disk()

    # 1) Excel download
    excel_bytes = get_excel_bytes(df)
    st.download_button(
        label="📥 단어장 다운로드 (Excel)",
        data=excel_bytes,
        file_name="vocab.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    # 2) Word(docs) download
    # word_bytes = get_word_bytes(df)
    # 2-1) Word(docs) 포맷에 따른 변경
    vocab = st.session_state["vocab"]
    word_bytes = get_word_bytes_from_vocab(vocab)
    st.download_button(
        label="📥 단어장 다운로드 (Word)",
        data=word_bytes,
        file_name="vocab.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # 3) PDF download
    # pdf_bytes = get_pdf_bytes(df)
    pdf_bytes = get_pdf_bytes_from_vocab(vocab)
    st.download_button(
        label="📥 단어장 다운로드 (PDF)",
        data=pdf_bytes,
        file_name="vocab.pdf",
        mime="application/pdf",
    )

    # 4) 발음 zip download 버튼
    audio_zip = get_audio_zip_from_vocab(vocab)
    st.download_button(
        label="🔊 단어/예문 발음 다운로드 (ZIP)",
        data=audio_zip,
        file_name="vocab_audio.zip",
        mime="application/zip",
    )
else:
    st.info("아직 저장된 단어가 없습니다.")
